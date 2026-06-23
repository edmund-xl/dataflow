from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from .artifacts import write_acceptance_checklist, write_findings_workbook, write_logic_check_results, write_validation_json
from .models import Finding, GraphModel, WorkbookData


def generate_reports(
    workbook: WorkbookData,
    graph: GraphModel,
    validation_findings: list[Finding],
    risk_findings: list[Finding],
    reports_dir: Path,
    env: str,
    version: str,
) -> None:
    reports_dir.mkdir(parents=True, exist_ok=True)
    write_findings_workbook(reports_dir / "validation_report.xlsx", "Validation", validation_findings)
    write_validation_json(validation_findings, reports_dir)
    write_findings_workbook(reports_dir / "issue_risk_register.xlsx", "Issues_Risks", validation_findings + risk_findings)
    write_logic_check_results(reports_dir / "logic_check_results.json", risk_findings)
    write_acceptance_checklist(reports_dir / "acceptance_checklist.xlsx", validation_findings, risk_findings, graph)
    _write_logic_mapping_docx(reports_dir / "logic_mapping_validation_report.docx", workbook, graph, validation_findings, risk_findings, env, version)


def _write_logic_mapping_docx(
    path: Path,
    workbook: WorkbookData,
    graph: GraphModel,
    validation_findings: list[Finding],
    risk_findings: list[Finding],
    env: str,
    version: str,
) -> None:
    doc = Document()
    generated_at = datetime.now(UTC).isoformat()
    conclusion = _review_conclusion(validation_findings, risk_findings)
    pending_count = len([f for f in validation_findings + risk_findings if f.status == "Pending_Confirmation"])
    p0_p1_validation = len([f for f in validation_findings if f.severity in {"P0", "P1"}])
    p0_risk = len([f for f in risk_findings if f.severity == "P0"])

    doc.add_heading("中文版本", level=0)
    doc.add_heading("Dataflow Project 逻辑映射验证报告", level=1)
    doc.add_heading("一、验收结论", level=2)
    doc.add_paragraph(f"结论：{conclusion}")
    doc.add_paragraph("判定口径：存在 P0 问题时为 Blocked；存在 P1 或待确认项时为 Needs Review；无阻断和待确认项时为 Pass。")
    doc.add_heading("摘要", level=2)
    doc.add_paragraph("本文记录本次数据采集包经过智能体处理后的逻辑映射验证结果。报告覆盖输入工作簿、图模型规模、校验结果、风险结果和再生成规则，用于评审数据流图是否具备交付条件。")
    doc.add_heading("关键词", level=2)
    doc.add_paragraph("逻辑映射；数据流图；图模型；校验报告；风险检查；可重复生成。")
    doc.add_heading("二、基础信息", level=2)
    _add_metric_table(
        doc,
        "指标",
        "数值",
        [
            ("环境", env),
            ("版本", version),
            ("生成时间", generated_at),
            ("输入工作簿", workbook.path.name),
        ],
    )
    doc.add_heading("三、验证结果摘要", level=2)
    _add_metric_table(
        doc,
        "指标",
        "数值",
        [
            ("节点数量", str(len(graph.nodes))),
            ("关系数量", str(len(graph.edges))),
            ("校验问题数量", str(len(validation_findings))),
            ("风险问题数量", str(len(risk_findings))),
            ("P0/P1 校验问题数量", str(p0_p1_validation)),
            ("P0 风险数量", str(p0_risk)),
            ("待确认问题数量", str(pending_count)),
        ],
    )
    doc.add_heading("四、校验问题", level=2)
    _add_finding_table_cn(doc, validation_findings)
    doc.add_heading("五、一致性与风险问题", level=2)
    _add_finding_table_cn(doc, risk_findings)
    doc.add_heading("六、开源授权", level=2)
    doc.add_paragraph("本项目采用 MIT License 开源授权。版权归属 edmund-xl；使用者可以在遵守许可证条款的前提下复制、使用、修改、分发和商用本软件副本。完整授权文本见交付包或仓库根目录中的 LICENSE 文件。")
    doc.add_heading("七、结论与再生成规则", level=2)
    doc.add_paragraph("如果生成图或报告存在错误，必须修正源工作簿并重新生成交付包，不允许手工修改生成产物。")

    doc.add_page_break()
    doc.add_heading("English Version", level=0)
    doc.add_heading("Dataflow Project Logic Mapping Validation Report", level=1)
    doc.add_heading("1. Acceptance Conclusion", level=2)
    doc.add_paragraph(f"Conclusion: {conclusion}")
    doc.add_paragraph("Decision rule: P0 findings mean Blocked; P1 or pending-confirmation findings mean Needs Review; no blocking or pending findings means Pass.")
    doc.add_heading("Abstract", level=2)
    doc.add_paragraph("This report records the logic-mapping validation results produced by the agent from the current Data Collection Package. It covers the input workbook, graph-model size, validation results, risk results, and regeneration rule, and is used to review whether the data-flow diagram is ready for delivery.")
    doc.add_heading("Keywords", level=2)
    doc.add_paragraph("Logic mapping; data flow diagram; graph model; validation report; risk check; reproducible generation.")
    doc.add_heading("2. Basic Information", level=2)
    _add_metric_table(
        doc,
        "Metric",
        "Value",
        [
            ("Environment", env),
            ("Version", version),
            ("Generated at", generated_at),
            ("Input workbook", workbook.path.name),
        ],
    )
    doc.add_heading("3. Validation Summary", level=2)
    _add_metric_table(
        doc,
        "Metric",
        "Value",
        [
            ("Graph nodes", str(len(graph.nodes))),
            ("Graph edges", str(len(graph.edges))),
            ("Validation findings", str(len(validation_findings))),
            ("Risk findings", str(len(risk_findings))),
            ("P0/P1 validation findings", str(p0_p1_validation)),
            ("P0 risk findings", str(p0_risk)),
            ("Pending confirmation findings", str(pending_count)),
        ],
    )
    doc.add_heading("4. Validation Findings", level=2)
    _add_finding_table(doc, validation_findings)
    doc.add_heading("5. Consistency And Risk Findings", level=2)
    _add_finding_table(doc, risk_findings)
    doc.add_heading("6. Open-Source License", level=2)
    doc.add_paragraph("This project is released under the MIT License. Copyright remains with edmund-xl, and users may copy, use, modify, distribute, and use the software commercially subject to the license terms. The full license text is available in the delivery package or the repository root LICENSE file.")
    doc.add_heading("7. Conclusion And Regeneration Rule", level=2)
    doc.add_paragraph("If a generated diagram or report is wrong, correct the source workbook and regenerate the package. Do not edit generated artifacts manually.")
    doc.save(path)


def _add_metric_table(doc: Document, left_label: str, right_label: str, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = left_label
    table.rows[0].cells[1].text = right_label
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value


def _review_conclusion(validation_findings: list[Finding], risk_findings: list[Finding]) -> str:
    findings = validation_findings + risk_findings
    if any(f.severity == "P0" for f in findings):
        return "Blocked"
    if any(f.severity == "P1" or f.status == "Pending_Confirmation" for f in findings):
        return "Needs Review"
    return "Pass"


def _add_finding_table_cn(doc: Document, findings: list[Finding]) -> None:
    if not findings:
        doc.add_paragraph("无问题。")
        return
    table = doc.add_table(rows=1, cols=5)
    for idx, label in enumerate(["门禁", "严重度", "工作表", "行", "问题摘要"]):
        table.rows[0].cells[idx].text = label
    for finding in findings:
        row = table.add_row().cells
        row[0].text = _gate_cn(finding.gate)
        row[1].text = finding.severity
        row[2].text = finding.sheet
        row[3].text = finding.row_id
        row[4].text = _finding_summary_cn(finding)


def _add_finding_table(doc: Document, findings: list[Finding]) -> None:
    if not findings:
        doc.add_paragraph("No findings.")
        return
    table = doc.add_table(rows=1, cols=5)
    for idx, label in enumerate(["Gate", "Severity", "Sheet", "Row", "Message"]):
        table.rows[0].cells[idx].text = label
    for finding in findings:
        row = table.add_row().cells
        row[0].text = finding.gate
        row[1].text = finding.severity
        row[2].text = finding.sheet
        row[3].text = finding.row_id
        row[4].text = finding.message


def _gate_cn(gate: str) -> str:
    return {
        "Gate 1": "门禁一",
        "Gate 2": "门禁二",
        "Gate 3": "门禁三",
        "Gate 4": "门禁四",
        "Gate 5": "门禁五",
    }.get(gate, gate)


def _finding_summary_cn(finding: Finding) -> str:
    if finding.status == "Pending_Confirmation":
        return "该记录仍处于待确认状态，不能作为最终事实直接验收。"
    if "monitoring" in finding.message.lower():
        return "关键对象缺少监控覆盖记录，需要补充监控、日志或告警信息。"
    if "firewall" in finding.message.lower():
        return "关键依赖缺少防火墙关联或放行原因，需要补充规则或例外说明。"
    if "foreign key" in finding.message.lower() or "does not exist" in finding.message.lower():
        return "引用对象不存在，需要修正标识或补充目标记录。"
    if "required field" in finding.message.lower() or "blank" in finding.message.lower():
        return "必填字段为空，需要在源工作簿中补充。"
    if "duplicate" in finding.message.lower():
        return "主键或记录重复，需要去重或调整标识。"
    return "智能体发现该记录需要复核，请结合机器可读报告查看原始细节。"
