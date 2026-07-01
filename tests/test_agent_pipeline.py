from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZipFile

import pytest
from docx import Document
from openpyxl import load_workbook

from dataflow_agent.architecture_findings import build_completeness_findings, write_architecture_findings
from dataflow_agent.constants import find_workbook
from dataflow_agent.diagram_renderer import VIEWS, render_diagrams, render_service_drilldown
from dataflow_agent.graph_builder import build_graph
from dataflow_agent.normalizer import normalize_workbook
from dataflow_agent.pipeline import run_all
from dataflow_agent.port_index import build_service_port_index
from dataflow_agent.risk_checker import check_risks
from dataflow_agent.schema import load_schema
from dataflow_agent.validator import validate_workbook
from dataflow_agent.xlsx_reader import read_workbook
from dataflow_agent.merge import merge_dcps


ROOT = Path(__file__).resolve().parents[1]
SAMPLE_DCP = ROOT / "samples" / "DCP_v0.1"
SAMPLE_WORKBOOK = SAMPLE_DCP / "dataflow_collection_template_v0.1.xlsx"
CLEAN_SAMPLE_DCP = ROOT / "samples" / "DCP_clean_v0.1"
CLEAN_SAMPLE_WORKBOOK = CLEAN_SAMPLE_DCP / "dataflow_collection_template_v0.1.xlsx"


def _main_dataflow_lines(svg: str) -> list[str]:
    lines = re.findall(r'<polyline[^>]+data-overview-role="main-dataflow"[^>]+', svg)
    return [line for line in lines if "main-dataflow-halo" not in line]


def _polyline_points(line: str) -> list[tuple[float, float]]:
    match = re.search(r'points="([^"]+)"', line)
    if not match:
        return []
    points: list[tuple[float, float]] = []
    for pair in match.group(1).split():
        x_value, y_value = pair.split(",", 1)
        points.append((float(x_value), float(y_value)))
    return points


def _script_env() -> dict[str, str]:
    import os

    env = dict(os.environ)
    env["DATAFLOW_PYTHON"] = sys.executable
    return env


def test_full_run_outputs_package(tmp_path: Path) -> None:
    state = run_all(CLEAN_SAMPLE_DCP, tmp_path, "testnetv2", "v0.1-demo")

    package_dir = tmp_path / "dataflow_package_v0.1-demo"
    assert (tmp_path / "dataflow_package_v0.1-demo.zip").exists()
    assert (package_dir / "normalized" / "nodes.csv").exists()
    assert (package_dir / "normalized" / "edges.csv").exists()
    assert (package_dir / "normalized" / "dataflow_graph.json").exists()
    assert (package_dir / "normalized" / "dataflow_graph.yaml").exists()
    assert (package_dir / "reports" / "validation_report.xlsx").exists()
    assert (package_dir / "reports" / "architecture_findings.md").exists()
    assert (package_dir / "reports" / "architecture_findings.json").exists()
    assert (package_dir / "reports" / "logic_mapping_validation_report.docx").exists()
    assert (package_dir / "reports" / "issue_risk_register.xlsx").exists()
    assert (package_dir / "reports" / "acceptance_checklist.xlsx").exists()
    assert (package_dir / "metadata.json").exists()
    assert state.zip_path and state.zip_path.exists()

    metadata = json.loads((package_dir / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["environment"] == "testnetv2"
    assert metadata["version"] == "v0.1-demo"
    assert metadata["schema_version"] == "workbook_schema.v0.1"
    assert metadata["template_version"] == "dataflow_template.v1.0"
    assert metadata["input_file_hash"]
    architecture_findings = (package_dir / "reports" / "architecture_findings.md").read_text(encoding="utf-8")
    assert "# 架构问题分析报告" in architecture_findings
    assert "# Architecture Findings Report" in architecture_findings
    assert "真实数据流链路" in architecture_findings
    assert "nonexistent relationships are not invented" in architecture_findings
    assert "总览图就绪度：READY" in architecture_findings
    architecture_json = json.loads((package_dir / "reports" / "architecture_findings.json").read_text(encoding="utf-8"))
    assert "coverage_matrix" in architecture_json
    assert "findings" in architecture_json
    assert architecture_json["conclusion"] == "PASS"
    assert any(observation["category"] == "Executive Overview" for observation in architecture_json["review_observations"])


def test_clean_sample_has_no_validation_or_risk_findings() -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(CLEAN_SAMPLE_WORKBOOK, schema), schema)
    validation = validate_workbook(workbook, schema)
    graph = build_graph(workbook)
    risks = check_risks(workbook, graph)

    assert validation.findings == []
    assert risks == []
    assert graph.dropped_edges == []


def test_sample_validates_and_builds_expected_edges() -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(SAMPLE_WORKBOOK, schema), schema)
    validation = validate_workbook(workbook, schema)
    graph = build_graph(workbook)

    assert validation.findings == []
    edge_types = {edge.type for edge in graph.edges}
    assert {"runs_on", "calls", "calls_external", "reads_from", "uses_sa", "monitored_by"}.issubset(edge_types)


def test_dropped_edges_and_edge_metadata_are_reported(tmp_path: Path) -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(SAMPLE_WORKBOOK, schema), schema)
    workbook.sheets["04_Services"][0]["Running_On_Instance_ID"] = "missing-instance"

    graph = build_graph(workbook)
    risks = check_risks(workbook, graph)

    assert graph.dropped_edges
    assert any(edge.target == "missing-instance" and "does not exist" in edge.reason for edge in graph.dropped_edges)
    assert any(f.gate == "Gate 5" and "Dropped graph edge" in f.message for f in risks)
    calls_edge = next(edge for edge in graph.edges if edge.type == "calls")
    assert calls_edge.metadata["source_sheet"] == "05_Dependencies"
    assert calls_edge.metadata["confirmation_status"] == "Confirmed"

    from dataflow_agent.artifacts import write_graph_artifacts

    write_graph_artifacts(graph, tmp_path)
    dropped_csv = (tmp_path / "dropped_edges.csv").read_text(encoding="utf-8")
    assert "missing-instance" in dropped_csv


def test_duplicate_primary_key_is_validation_error() -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(SAMPLE_WORKBOOK, schema), schema)
    workbook.sheets["04_Services"].append(dict(workbook.sheets["04_Services"][0]))

    validation = validate_workbook(workbook, schema)

    assert any(f.sheet == "04_Services" and "Duplicate primary key" in f.message for f in validation.findings)


def test_broken_foreign_key_is_validation_error() -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(SAMPLE_WORKBOOK, schema), schema)
    workbook.sheets["04_Services"][0]["Running_On_Instance_ID"] = "missing-instance"

    validation = validate_workbook(workbook, schema)

    assert any(f.gate == "Gate 2" and f.field == "Running_On_Instance_ID" and f.severity == "P0" for f in validation.findings)


def test_rejected_rows_do_not_enter_graph_and_pending_rows_enter_risk_register() -> None:
    schema = load_schema()
    workbook = read_workbook(SAMPLE_WORKBOOK, schema)
    workbook.sheets["04_Services"][0]["Confirmation_Status"] = "Rejected"
    normalized = normalize_workbook(workbook, schema)
    graph = build_graph(normalized)
    risks = check_risks(normalized, graph)

    assert "svc-nginx-entry" not in graph.nodes
    assert any(f.status == "Pending_Confirmation" for f in risks)


def test_gcp_risk_rules_cover_security_network_iam_and_monitoring() -> None:
    schema = load_schema()
    workbook = read_workbook(SAMPLE_WORKBOOK, schema)
    workbook.sheets["08_Cloud_Armor"] = []
    for row in workbook.sheets["02_Networks"]:
        row["NAT_Name"] = ""
        row["PSC_or_Peering_Name"] = ""
    workbook.sheets["07_Firewalls"][0]["Ports"] = "9443"
    workbook.sheets["09_IAM_SA"].append(
        {
            **workbook.sheets["09_IAM_SA"][0],
            "Record_ID": "rec-test-admin",
            "IAM_Binding_ID": "iam-test-admin",
            "Service_Account_ID": "sa-admin",
            "Used_By_Service_ID": "svc-rpc-api",
            "Role": "roles/owner",
            "Justification": "",
            "Is_High_Privilege": "Yes",
        }
    )
    workbook.sheets["10_Monitoring"] = []
    normalized = normalize_workbook(workbook, schema)
    graph = build_graph(normalized)

    risks = check_risks(normalized, graph)
    messages = "\n".join(f.message for f in risks)

    assert "no Cloud Armor/LB/nginx protection row" in messages
    assert "no NAT/egress record" in messages
    assert "no PSC/VPC Peering record" in messages
    assert "does not match dependency" in messages
    assert "High privilege IAM binding has no justification" in messages
    assert "P0 service" in messages
    assert "Critical dependency" in messages
    assert "Sensitive data asset" in messages


def test_all_diagram_views_render_nonempty_files(tmp_path: Path) -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(SAMPLE_WORKBOOK, schema), schema)
    graph = build_graph(workbook)

    outputs = render_diagrams(graph, tmp_path)

    expected_count = len(VIEWS) * 6
    assert len(outputs) == expected_count
    for path in outputs:
        assert path.exists()
        assert path.stat().st_size > 100

    overview_svg = (tmp_path / "00_overview.svg").read_text(encoding="utf-8")
    service_svg = (tmp_path / "03_service_dependency_layer.svg").read_text(encoding="utf-8")
    security_svg = (tmp_path / "05_security_monitoring_layer.svg").read_text(encoding="utf-8")
    overview_drawio = ElementTree.parse(tmp_path / "00_overview.drawio")
    overview_graphml = ElementTree.parse(tmp_path / "00_overview.graphml")
    legacy = "mega" + "eth"
    for svg in (overview_svg, service_svg):
        assert "Entry context / perimeter control" in svg
        assert "Primary graph dataflow" in svg
        assert "Controls and runtime summary" in svg
        assert "Graph edge ledger" in svg
        assert "Semantic guardrail" in svg
        assert 'data-overview-role="main-dataflow-halo"' in svg
        assert 'data-overview-role="main-dataflow"' in svg
        assert 'data-edge-number="E' in svg
    assert "Subnet" not in overview_svg
    assert 'data-edge-type="allowed_by"' not in _main_dataflow_lines(service_svg)
    assert 'data-edge-type="uses_runtime"' not in _main_dataflow_lines(service_svg)
    assert "Firewall Rule" not in service_svg
    assert "Structurizr/C4-style architecture view" in security_svg
    assert "#101827" not in security_svg
    assert 'data-security-row="' not in security_svg
    drawio_edges = [cell for cell in overview_drawio.findall(".//mxCell") if cell.attrib.get("edge") == "1"]
    assert drawio_edges
    assert all(cell.attrib.get("graphEdgeId") for cell in drawio_edges)
    assert not any(cell.attrib.get("sourceNodeId") == "prod-lb-public" and cell.attrib.get("targetNodeId") == "svc-nginx-entry" for cell in drawio_edges)
    assert overview_graphml.findall(".//{http://graphml.graphdrawing.org/xmlns}edge")
    assert legacy not in overview_svg.lower()
    assert legacy not in security_svg.lower()


def test_overview_renderer_uses_graph_truthful_edges(tmp_path: Path) -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(SAMPLE_WORKBOOK, schema), schema)
    graph = build_graph(workbook)

    render_diagrams(graph, tmp_path)

    overview_svg = (tmp_path / "00_overview.svg").read_text(encoding="utf-8")
    main_lines = _main_dataflow_lines(overview_svg)
    main_badges = re.findall(r'data-edge-badge-id="[^"]+"', overview_svg)
    compact_badges = re.findall(r'<rect[^>]+width="46"[^>]+data-overview-role="main-dataflow-label"[^>]+', overview_svg)
    rendered_edge_ids = {match.group(1) for line in main_lines if (match := re.search(r'data-edge-id="([^"]+)"', line))}
    rendered_edge_types = {match.group(1) for line in main_lines if (match := re.search(r'data-edge-type="([^"]+)"', line))}
    rendered_routes = {match.group(1) for line in main_lines if (match := re.search(r'points="([^"]+)"', line))}
    terminal_bus_y = []
    for line in main_lines:
        edge_type = re.search(r'data-edge-type="([^"]+)"', line)
        if not edge_type or edge_type.group(1) not in {"calls_external", "reads_from", "writes_to"}:
            continue
        points = _polyline_points(line)
        horizontal_segments = [
            first[1]
            for first, second in zip(points, points[1:])
            if abs(first[1] - second[1]) < 0.01 and abs(first[0] - second[0]) > 140
        ]
        assert horizontal_segments
        terminal_bus_y.append(round(horizontal_segments[0], 1))

    assert rendered_edge_ids >= {"edge-2000", "edge-2001", "edge-2002", "edge-2003", "edge-2004", "edge-2005"}
    assert len(main_badges) == len(main_lines)
    assert len(compact_badges) == len(main_lines)
    assert "HTTP/JSON 8545" in overview_svg
    assert "TCP 9090" in overview_svg
    assert len(rendered_routes) == len(main_lines)
    assert len(terminal_bus_y) == len(set(terminal_bus_y))
    assert rendered_edge_types <= {"calls", "calls_external", "reads_from", "writes_to", "depends_on"}
    assert not (rendered_edge_types & {"runs_on", "runs_on_runtime", "allowed_by", "uses_sa", "monitored_by", "protected_by"})
    assert 'data-source="prod-lb-public" data-target="svc-nginx-entry"' not in overview_svg
    assert 'data-edge-type="allowed_by"' not in "\n".join(main_lines)
    assert 'data-edge-type="uses_sa"' not in "\n".join(main_lines)


def test_diagrams_show_non_final_statuses(tmp_path: Path) -> None:
    schema = load_schema()
    workbook = read_workbook(SAMPLE_WORKBOOK, schema)
    workbook.sheets["04_Services"][0]["Confirmation_Status"] = "Pending_Confirmation"
    workbook.sheets["04_Services"][1]["Confirmation_Status"] = "Accepted_Exception"
    workbook.sheets["04_Services"][2]["Confirmation_Status"] = "Auto_Detected"
    normalized = normalize_workbook(workbook, schema)
    graph = build_graph(normalized)

    render_diagrams(graph, tmp_path)

    overview_svg = (tmp_path / "00_overview.svg").read_text(encoding="utf-8")
    overview_mmd = (tmp_path / "00_overview.mmd").read_text(encoding="utf-8")
    assert "PENDING" in overview_svg
    assert "EXC" in overview_svg
    assert "AUTO" in overview_svg
    assert 'aria-label="' in overview_svg
    assert 'data-risk-level="review"' in overview_svg
    assert "<title>" in overview_svg
    assert "Pending_Confirmation" in overview_mmd
    assert "Accepted_Exception" in overview_mmd


def test_k8s_service_with_data_asset_dependency_end_to_end(tmp_path: Path) -> None:
    schema = load_schema()
    workbook = read_workbook(SAMPLE_WORKBOOK, schema)
    for field in ["Runtime_Type", "Runtime_ID", "Runtime_Name", "Runtime_Namespace", "Runtime_Cluster", "Runtime_Region"]:
        assert field in workbook.headers["04_Services"]
    for field in ["Target_Type", "Target_ID", "Interaction_Mode"]:
        assert field in workbook.headers["05_Dependencies"]
    service = next(row for row in workbook.sheets["04_Services"] if row["Service_ID"] == "svc-rpc-api")
    service.update(
        {
            "Runtime_Type": "Kubernetes",
            "Runtime_ID": "k8s-deploy-rpc",
            "Runtime_Name": "rpc-api-deployment",
            "Runtime_Namespace": "rpc",
            "Runtime_Cluster": "cluster-main",
            "Runtime_Region": "us-central1",
        }
    )
    dep = next(row for row in workbook.sheets["05_Dependencies"] if row["Dependency_ID"] == "dep-rpc-sequencer")
    dep.update(
        {
            "Target_Service_ID": "",
            "Target_Data_Asset_ID": "",
            "Target_Type": "data_asset",
            "Target_ID": "data-cloudsql-state",
            "Interaction_Mode": "write",
            "Direction": "write",
        }
    )
    normalized = normalize_workbook(workbook, schema)

    validation = validate_workbook(normalized, schema)
    graph = build_graph(normalized)
    risks = check_risks(normalized, graph)
    render_service_drilldown(graph, "svc-rpc-api", tmp_path)

    assert not [f for f in validation.findings if f.severity in {"P0", "P1"}]
    assert "k8s-deploy-rpc" in graph.nodes
    runtime_edge = next(edge for edge in graph.edges if edge.type == "runs_on_runtime" and edge.source == "svc-rpc-api")
    assert runtime_edge.metadata["runtime_context"]["cluster"] == "cluster-main"
    assert any(edge.type == "writes_to" and edge.source == "svc-rpc-api" and edge.target == "data-cloudsql-state" for edge in graph.edges)
    assert not [f for f in risks if f.gate == "Gate 5"]
    drilldown_svg = (tmp_path / "service_drilldown_svc-rpc-api.svg").read_text(encoding="utf-8")
    assert "rpc-api-deployment" in drilldown_svg


def test_sample_workbook_exposes_runtime_and_target_fields() -> None:
    schema = load_schema()
    workbook = read_workbook(SAMPLE_WORKBOOK, schema)

    service = next(row for row in workbook.sheets["04_Services"] if row["Service_ID"] == "svc-rpc-api")
    dependency = next(row for row in workbook.sheets["05_Dependencies"] if row["Dependency_ID"] == "dep-rpc-sequencer")

    for field in ["Runtime_Type", "Runtime_ID", "Runtime_Name", "Runtime_Namespace", "Runtime_Cluster", "Runtime_Region"]:
        assert field in workbook.headers["04_Services"]
        assert field in service
    for field in ["Target_Type", "Target_ID", "Interaction_Mode"]:
        assert field in workbook.headers["05_Dependencies"]
        assert field in dependency
    assert service["Runtime_Type"] == "Kubernetes"
    assert dependency["Target_Type"] == "service"


def test_service_drilldown_renders_expected_artifacts(tmp_path: Path) -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(SAMPLE_WORKBOOK, schema), schema)
    graph = build_graph(workbook)

    outputs = render_service_drilldown(graph, "svc-rpc-api", tmp_path)

    assert len(outputs) == 6
    for path in outputs:
        assert path.exists()
        assert path.stat().st_size > 100
    svg = (tmp_path / "service_drilldown_svc-rpc-api.svg").read_text(encoding="utf-8")
    drawio = ElementTree.parse(tmp_path / "service_drilldown_svc-rpc-api.drawio")
    graphml = ElementTree.parse(tmp_path / "service_drilldown_svc-rpc-api.graphml")
    assert "Service Drilldown: svc-rpc-api" in svg
    assert "RPC API" in svg
    assert drawio.findall(".//mxCell")
    assert graphml.findall(".//{http://graphml.graphdrawing.org/xmlns}edge")


def test_service_drilldown_supports_depth_direction_theme_and_risk_focus(tmp_path: Path) -> None:
    schema = load_schema()
    workbook = read_workbook(SAMPLE_WORKBOOK, schema)
    next(row for row in workbook.sheets["05_Dependencies"] if row["Dependency_ID"] == "dep-rpc-sequencer")["Confirmation_Status"] = "Pending_Confirmation"
    normalized = normalize_workbook(workbook, schema)
    graph = build_graph(normalized)

    outputs = render_service_drilldown(graph, "svc-rpc-api", tmp_path, depth=2, direction="downstream", theme="dark", risk_focus=True)

    assert len(outputs) == 6
    svg = (tmp_path / "service_drilldown_svc-rpc-api.svg").read_text(encoding="utf-8")
    assert "depth=2" in svg
    assert "#101827" in svg
    assert 'data-risk-level="review"' in svg


def test_cli_input_can_be_copied_to_fresh_dcp(tmp_path: Path) -> None:
    dcp = tmp_path / "DCP_v0.1"
    dcp.mkdir()
    shutil.copy2(SAMPLE_WORKBOOK, dcp / SAMPLE_WORKBOOK.name)

    state = run_all(dcp, tmp_path / "out", "testnetv2", "v0.1-demo")

    assert state.validation.findings == []
    assert state.zip_path and state.zip_path.exists()


def test_merge_identical_dcp_deduplicates_rows(tmp_path: Path) -> None:
    result = merge_dcps([SAMPLE_DCP, SAMPLE_DCP], tmp_path, "v0.1-demo")

    assert result.workbook_path.exists()
    assert (result.merged_dcp / "merge_report.xlsx").exists()
    assert (result.merged_dcp / "merge_report.json").exists()
    assert (result.merged_dcp / "merge_lineage.json").exists()
    assert result.lineage
    assert result.duplicate_count > 0
    assert result.conflict_count == 0


def test_merge_conflicts_block_final_package_unless_allowed(tmp_path: Path) -> None:
    first = tmp_path / "first"
    second = tmp_path / "second"
    first.mkdir()
    second.mkdir()
    shutil.copy2(SAMPLE_WORKBOOK, first / SAMPLE_WORKBOOK.name)
    second_workbook = second / SAMPLE_WORKBOOK.name
    shutil.copy2(SAMPLE_WORKBOOK, second_workbook)
    wb = load_workbook(second_workbook)
    ws = wb["04_Services"]
    header_row = next(row for row in range(1, ws.max_row + 1) if any(cell.value == "Service_ID" for cell in ws[row]))
    headers = [cell.value for cell in ws[header_row]]
    service_id_col = headers.index("Service_ID") + 1
    service_name_col = headers.index("Service_Name") + 1
    for row in range(header_row + 1, ws.max_row + 1):
        if ws.cell(row, service_id_col).value == "svc-rpc-api":
            ws.cell(row, service_name_col).value = "RPC API Conflict"
            break
    wb.save(second_workbook)

    blocked_out = tmp_path / "blocked"
    blocked = subprocess.run(
        [sys.executable, "-m", "dataflow_agent.cli", "merge", str(first), str(second), "--output", str(blocked_out)],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    assert blocked.returncode == 1
    assert "Resolve conflicts" in blocked.stderr
    assert not list(blocked_out.glob("dataflow_package_*.zip"))

    draft_out = tmp_path / "draft"
    subprocess.run(
        [sys.executable, "-m", "dataflow_agent.cli", "merge", str(first), str(second), "--output", str(draft_out), "--allow-conflicts"],
        cwd=ROOT,
        check=True,
    )
    metadata = json.loads((draft_out / "dataflow_package_v0.1-demo" / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["delivery_status"] == "Draft"
    assert metadata["merge_conflict_count"] > 0
    assert (draft_out / "dataflow_package_v0.1-demo" / "reports" / "merge_lineage.json").exists()
    assert (draft_out / "dataflow_package_v0.1-demo" / "reports" / "conflict_diff.xlsx").exists()
    draft_conflicts = (draft_out / "dataflow_package_v0.1-demo" / "reports" / "DRAFT_CONFLICTS.md").read_text(encoding="utf-8")
    assert "first-kept" in draft_conflicts
    assert "Service_Name" in draft_conflicts
    assert "RPC API Conflict" in draft_conflicts


def test_merge_conflict_count_decreases_after_source_fix(tmp_path: Path) -> None:
    first = tmp_path / "first"
    second = tmp_path / "second"
    first.mkdir()
    second.mkdir()
    shutil.copy2(SAMPLE_WORKBOOK, first / SAMPLE_WORKBOOK.name)
    second_workbook = second / SAMPLE_WORKBOOK.name
    shutil.copy2(SAMPLE_WORKBOOK, second_workbook)
    wb = load_workbook(second_workbook)
    ws = wb["04_Services"]
    header_row = next(row for row in range(1, ws.max_row + 1) if any(cell.value == "Service_ID" for cell in ws[row]))
    headers = [cell.value for cell in ws[header_row]]
    service_id_col = headers.index("Service_ID") + 1
    service_name_col = headers.index("Service_Name") + 1
    target_row = None
    for row in range(header_row + 1, ws.max_row + 1):
        if ws.cell(row, service_id_col).value == "svc-rpc-api":
            target_row = row
            ws.cell(row, service_name_col).value = "RPC API Conflict"
            break
    assert target_row is not None
    wb.save(second_workbook)

    conflicted = merge_dcps([first, second], tmp_path / "conflicted", "v0.1-demo")
    assert conflicted.conflict_count > 0
    assert conflicted.conflict_diffs

    ws.cell(target_row, service_name_col).value = "RPC API"
    wb.save(second_workbook)
    resolved = merge_dcps([first, second], tmp_path / "resolved", "v0.1-demo")
    assert resolved.conflict_count < conflicted.conflict_count


def test_script_check_dcp_runs_with_defaults(tmp_path: Path) -> None:
    subprocess.run(["scripts/check_dcp.sh", "samples/DCP_clean_v0.1"], cwd=ROOT, check=True, env=_script_env())

    check_dir = CLEAN_SAMPLE_DCP / "agent_check"
    assert (check_dir / "check_summary.md").exists()
    assert (check_dir / "fix_list.md").exists()
    assert (check_dir / "architecture_findings.md").exists()
    assert (check_dir / "architecture_findings.json").exists()
    assert "自检状态：`PASS`" in (check_dir / "check_summary.md").read_text(encoding="utf-8")
    assert "分析结论：`PASS`" in (check_dir / "architecture_findings.md").read_text(encoding="utf-8")
    assert "架构完整性发现：P0=0，P1=0，P2=0" in (check_dir / "check_summary.md").read_text(encoding="utf-8")

    custom_check = tmp_path / "agent_check"
    subprocess.run(["scripts/check_dcp.sh", "samples/DCP_clean_v0.1", "--output", str(custom_check)], cwd=ROOT, check=True, env=_script_env())
    assert (custom_check / "check_summary.md").exists()
    assert (custom_check / "fix_list.md").exists()
    assert (custom_check / "architecture_findings.md").exists()
    assert (custom_check / "architecture_findings.json").exists()


def test_script_build_package_runs_with_defaults(tmp_path: Path) -> None:
    subprocess.run(["scripts/build_dataflow_package.sh", "samples/DCP_clean_v0.1"], cwd=ROOT, check=True, env=_script_env())

    dist_dir = CLEAN_SAMPLE_DCP / "dist"
    assert any(path.name.startswith("dataflow_package_") and path.suffix == ".zip" for path in dist_dir.iterdir())

    custom_dist = tmp_path / "dist"
    subprocess.run(["scripts/build_dataflow_package.sh", "samples/DCP_clean_v0.1", "--output", str(custom_dist)], cwd=ROOT, check=True, env=_script_env())
    assert any(path.name.startswith("dataflow_package_") and path.suffix == ".zip" for path in custom_dist.iterdir())


def test_script_merge_dcp_runs_with_defaults(tmp_path: Path) -> None:
    subprocess.run(["scripts/merge_dcp.sh", "samples/DCP_clean_v0.1", "samples/DCP_clean_v0.1"], cwd=ROOT, check=True, env=_script_env())

    merge_reports = list((ROOT / "dist").glob("merged_dcp_*/merge_report.xlsx"))
    package_zips = list((ROOT / "dist").glob("dataflow_package_*.zip"))
    assert merge_reports
    assert package_zips

    custom_merge = tmp_path / "merge"
    subprocess.run(["scripts/merge_dcp.sh", "samples/DCP_clean_v0.1", "samples/DCP_clean_v0.1", "--output", str(custom_merge)], cwd=ROOT, check=True, env=_script_env())
    assert list(custom_merge.glob("merged_dcp_*/merge_report.xlsx"))
    assert list(custom_merge.glob("dataflow_package_*.zip"))


def test_script_service_drilldown_runs_with_defaults() -> None:
    subprocess.run(["scripts/build_service_drilldown.sh", "samples/DCP_clean_v0.1", "svc-rpc-api"], cwd=ROOT, check=True, env=_script_env())
    subprocess.run(["scripts/build_service_drilldown.sh", "samples/DCP_clean_v0.1", "svc-rpc-api", "--depth", "2", "--direction", "downstream", "--theme", "dark", "--risk-focus"], cwd=ROOT, check=True, env=_script_env())

    output_dir = CLEAN_SAMPLE_DCP / "dist" / "service_drilldown_svc-rpc-api"
    assert (output_dir / "service_drilldown_svc-rpc-api.svg").exists()
    assert (output_dir / "service_drilldown_svc-rpc-api.png").exists()
    assert (output_dir / "service_drilldown_svc-rpc-api.pdf").exists()
    assert (output_dir / "service_drilldown_svc-rpc-api.mmd").exists()
    assert (output_dir / "service_drilldown_svc-rpc-api.drawio").exists()
    assert (output_dir / "service_drilldown_svc-rpc-api.graphml").exists()


def test_service_port_query_runs_with_defaults(tmp_path: Path) -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(SAMPLE_WORKBOOK, schema), schema)
    index = build_service_port_index(workbook, "svc-rpc-api")
    assert "8545" in index["listen_ports"]
    assert index["outbound_dependencies"]

    subprocess.run(["scripts/query_service_ports.sh", "samples/DCP_clean_v0.1", "svc-rpc-api"], cwd=ROOT, check=True, env=_script_env())
    output_file = CLEAN_SAMPLE_DCP / "dist" / "service_ports_svc-rpc-api.json"
    assert output_file.exists()
    output = json.loads(output_file.read_text(encoding="utf-8"))
    assert output["service_id"] == "svc-rpc-api"
    assert output["source_row_id"] == "SVC-002"
    assert output["outbound_dependencies"][0]["graph_edges"][0]["edge_id"]

    custom_output = tmp_path / "custom_service_ports.json"
    subprocess.run(
        ["scripts/query_service_ports.sh", "samples/DCP_clean_v0.1", "svc-rpc-api", "--output", str(custom_output)],
        cwd=ROOT,
        check=True,
        env=_script_env(),
    )
    custom = json.loads(custom_output.read_text(encoding="utf-8"))
    assert custom["service_id"] == "svc-rpc-api"


def test_doctor_script_checks_environment() -> None:
    result = subprocess.run(["scripts/doctor.sh", "samples/DCP_clean_v0.1"], cwd=ROOT, check=True, env=_script_env(), text=True, capture_output=True)

    assert "Dataflow Agent Doctor" in result.stdout
    assert "READY" in result.stdout
    assert "WARN" in result.stdout
    assert "MISSING" in result.stdout
    assert "missing=0" in result.stdout


def test_scripts_do_not_embed_personal_python_paths() -> None:
    for path in [
        ROOT / "scripts" / "doctor.sh",
        ROOT / "scripts" / "check_dcp.sh",
        ROOT / "scripts" / "build_dataflow_package.sh",
        ROOT / "scripts" / "merge_dcp.sh",
        ROOT / "scripts" / "build_service_drilldown.sh",
        ROOT / "scripts" / "query_service_ports.sh",
        ROOT / "scripts" / "setup_env.sh",
    ]:
        text = path.read_text(encoding="utf-8")
        assert "DATAFLOW_PYTHON" in text
        assert "codex-runtimes" not in text
        assert "/Users/" not in text


def test_find_workbook_rejects_ambiguous_nonstandard_xlsx(tmp_path: Path) -> None:
    dcp = tmp_path / "ambiguous_dcp"
    dcp.mkdir()
    first = dcp / "first.xlsx"
    second = dcp / "second.xlsx"
    shutil.copy2(SAMPLE_WORKBOOK, first)
    shutil.copy2(SAMPLE_WORKBOOK, second)

    with pytest.raises(FileNotFoundError, match="Multiple workbook files found"):
        find_workbook(dcp)

    assert find_workbook(first) == first


def test_repository_docs_are_chinese_first_then_english() -> None:
    docs = [
        ROOT / "README.md",
        ROOT / "CHANGELOG.md",
        ROOT / "SECURITY.md",
        ROOT / "DATA_HANDLING.md",
        ROOT / "samples" / "README.md",
        ROOT / "docs" / "collector_quick_check_guide.md",
        ROOT / "docs" / "aggregation_operator_guide.md",
        ROOT / "docs" / "dataflow_agent_input_contract_v0.1.md",
        ROOT / "docs" / "devops_dcp_collection_manual.md",
        ROOT / "docs" / "devops_collection_filling_guide.md",
        ROOT / "docs" / "dcp_self_check_guide.md",
        ROOT / "docs" / "package_generation_guide.md",
    ]
    for path in docs:
        text = path.read_text(encoding="utf-8")
        assert "# 中文版本" in text
        assert "# English Version" in text
        assert text.index("# 中文版本") < text.index("# English Version")


def test_security_docs_and_sensitive_scan_exist() -> None:
    security = (ROOT / "SECURITY.md").read_text(encoding="utf-8")
    data_handling = (ROOT / "DATA_HANDLING.md").read_text(encoding="utf-8")
    samples = (ROOT / "samples" / "README.md").read_text(encoding="utf-8")

    assert "真实 DCP" in security
    assert "不得提交 Git" in data_handling
    assert "脱敏演示数据" in samples
    subprocess.run(["scripts/scan_sensitive.sh"], cwd=ROOT, check=True)


def test_github_actions_ci_covers_core_flow() -> None:
    workflow = (ROOT / ".github" / "workflows" / "ci.yml").read_text(encoding="utf-8")

    assert "python -m pytest -q" in workflow
    assert "scripts/scan_sensitive.sh" in workflow
    assert "scripts/doctor.sh samples/DCP_clean_v0.1" in workflow
    assert "dataflow-agent check samples/DCP_clean_v0.1" in workflow
    assert "dataflow-agent quick-build samples/DCP_clean_v0.1" in workflow
    assert "dataflow-agent merge samples/DCP_clean_v0.1 samples/DCP_clean_v0.1" in workflow
    assert "dataflow-agent drilldown --input samples/DCP_clean_v0.1 --service-id svc-rpc-api" in workflow
    assert "dataflow-agent query-port --input samples/DCP_clean_v0.1 --service-id svc-rpc-api" in workflow
    assert "actions/upload-artifact@v4" in workflow


def test_generated_docs_are_chinese_first_then_english(tmp_path: Path) -> None:
    state = run_all(SAMPLE_DCP, tmp_path, "testnetv2", "v0.1-demo")
    package_dir = tmp_path / "dataflow_package_v0.1-demo"
    package_readme = (package_dir / "README.md").read_text(encoding="utf-8")
    assert package_readme.index("# 中文版本") < package_readme.index("# English Version")

    doc = Document(package_dir / "reports" / "logic_mapping_validation_report.docx")
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_values = [cell.text for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()]
    assert paragraphs[0] == "中文版本"
    assert any(text in paragraphs for text in ["结论：Pass", "结论：Needs Review", "结论：Blocked"])
    assert "P0/P1 validation findings" in table_values
    assert "English Version" in paragraphs
    assert paragraphs.index("中文版本") < paragraphs.index("English Version")

    issue_register = load_workbook(package_dir / "reports" / "issue_risk_register.xlsx")
    headers = [cell.value for cell in issue_register.active[1]]
    for field in ["Owner", "Due_Date", "Exception_Decision", "Evidence_ID"]:
        assert field in headers
    architecture_findings = (package_dir / "reports" / "architecture_findings.md").read_text(encoding="utf-8")
    assert architecture_findings.index("# 中文版本") < architecture_findings.index("# English Version")
    assert "| Edge_ID | 类型 | 来源 | 目标 | 状态 | 来源记录 | Evidence_ID |" in architecture_findings
    assert "不存在的关系不会被补画或补写" in architecture_findings


def test_architecture_findings_include_review_observations(tmp_path: Path) -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(CLEAN_SAMPLE_WORKBOOK, schema), schema)
    workbook.sheets["09_IAM_SA"] = []

    firewall = workbook.sheets["07_Firewalls"][0]
    firewall["Firewall_ID"] = "fw-wide-admin"
    firewall["Direction"] = "ingress"
    firewall["Source_Allowed"] = "0.0.0.0/0"
    firewall["Ports"] = "22"
    firewall["Reason"] = "SSH admin temporary exception"
    firewall["Confirmation_Status"] = "Accepted_Exception"

    for row in workbook.sheets["10_Monitoring"]:
        row["Coverage_Status"] = "Partial"
        row["Dashboard_URL"] = ""
        row["Alert_Rule"] = "Not confirmed in repo"
        row["XDR_Coverage"] = "Unknown"

    graph = build_graph(workbook)
    output = tmp_path / "architecture_findings.md"
    write_architecture_findings(output, workbook, graph, [], [])
    content = output.read_text(encoding="utf-8")

    assert "## 审查观察项" in content
    assert "采集库存" in content
    assert "服务共" in content
    assert "依赖共" in content
    assert "数据资产共" in content
    assert "外部系统共" in content
    assert "网络记录共" in content
    assert "CI/CD 记录共" in content
    assert "Evidence 记录共" in content
    assert "P0 服务共" in content
    assert "fw-wide-admin" in content
    assert "0.0.0.0/0" in content
    assert "Partial=" in content
    assert "Alert_Rule 未确认" in content
    assert "当前没有 IAM / Service Account 记录" in content
    assert "No IAM / Service Account records are present" in content


def test_completeness_rule_matrix_detects_missing_cross_domain_information() -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(CLEAN_SAMPLE_WORKBOOK, schema), schema)

    service = next(row for row in workbook.sheets["04_Services"] if row["Service_ID"] == "svc-rpc-api")
    service["Service_Owner"] = ""
    service["Listen_Ports"] = ""

    dependency = next(row for row in workbook.sheets["05_Dependencies"] if row["Dependency_ID"] == "dep-rpc-sequencer")
    dependency["Auth_Method"] = ""
    workbook.sheets["07_Firewalls"] = [row for row in workbook.sheets["07_Firewalls"] if row.get("Related_Dependency_ID") != "dep-rpc-sequencer"]
    workbook.sheets["10_Monitoring"] = [row for row in workbook.sheets["10_Monitoring"] if row.get("Object_ID") != "dep-rpc-sequencer"]

    asset = next(row for row in workbook.sheets["06_Data_Assets"] if row["Data_Asset_ID"] == "data-cloudsql-state")
    asset["Backup_Policy"] = ""
    asset["Used_By_Service_ID"] = ""

    external = next(row for row in workbook.sheets["12_External_Services"] if row["External_ID"] == "ext-eigenda")
    external["Auth_Method"] = ""
    external["Purpose"] = ""
    external["Data_Classification"] = ""

    firewall = workbook.sheets["07_Firewalls"][0]
    firewall["Firewall_ID"] = "fw-wide-no-exception"
    firewall["Direction"] = "ingress"
    firewall["Source_Allowed"] = "0.0.0.0/0"
    firewall["Confirmation_Status"] = "Confirmed"

    iam = workbook.sheets["09_IAM_SA"][0]
    iam["Role"] = "roles/owner"
    iam["Is_High_Privilege"] = "Yes"
    iam["Justification"] = ""

    cicd = workbook.sheets["11_CICD"][0]
    cicd["Approval_Required"] = "No"
    cicd["Deployment_Account"] = ""
    cicd["Target_Service_ID"] = ""
    cicd["Target_Instance_ID"] = ""

    evidence = workbook.sheets["14_Evidence_Index"][0]
    evidence["Integrity_Note"] = ""

    graph = build_graph(workbook)
    findings = build_completeness_findings(workbook, graph)
    messages = "\n".join(f"{finding.severity} {finding.category} {finding.message_en}" for finding in findings)

    assert "Critical service svc-rpc-api has no Service_Owner" in messages
    assert "Critical service svc-rpc-api has no Listen_Ports" in messages
    assert "Critical dependency dep-rpc-sequencer has no related firewall rule" in messages
    assert "Dependency dep-rpc-sequencer has no Auth_Method" in messages
    assert "Critical dependency dep-rpc-sequencer has no direct monitoring coverage" in messages
    assert "Data asset data-cloudsql-state has no Backup_Policy" in messages
    assert "Data asset data-cloudsql-state has no Used_By_Service_ID" in messages
    assert "External service ext-eigenda has no Auth_Method" in messages
    assert "External service ext-eigenda has no Purpose" in messages
    assert "External service ext-eigenda has no Data_Classification" in messages
    assert "allows ingress from 0.0.0.0/0 without Accepted_Exception status" in messages
    assert "High privilege IAM binding" in messages
    assert "does not show required approval" in messages
    assert "has no Deployment_Account" in messages
    assert "has no deployment target" in messages
    assert "has no Integrity_Note" in messages


def test_overview_readiness_detects_missing_executive_summary_inputs() -> None:
    schema = load_schema()
    workbook = normalize_workbook(read_workbook(CLEAN_SAMPLE_WORKBOOK, schema), schema)
    workbook.sheets["08_Cloud_Armor"] = []
    workbook.sheets["09_IAM_SA"] = []
    for row in workbook.sheets["03_Servers"]:
        row["IP_External"] = ""
        row["Server_Role"] = "internal node"
        row["Network_Tag"] = "internal"
    for row in workbook.sheets["04_Services"]:
        row["Service_Name"] = row.get("Service_ID", "")
        row["Service_Role"] = "internal service"
        row["Description"] = ""
    for row in workbook.sheets["10_Monitoring"]:
        row["Coverage_Status"] = "Partial"

    graph = build_graph(workbook)
    findings = build_completeness_findings(workbook, graph)
    overview_findings = [finding for finding in findings if finding.category == "Executive Overview"]

    assert len(overview_findings) == 1
    assert overview_findings[0].severity == "P2"
    assert "Executive overview is not ready" in overview_findings[0].message_en
    assert "09_IAM_SA" in overview_findings[0].message_cn
    assert "10_Monitoring.Coverage_Status" in overview_findings[0].message_cn


def test_devops_docs_and_deterministic_agent_boundary_are_documented() -> None:
    readme = (ROOT / "README.md").read_text(encoding="utf-8")
    collection_guide = (ROOT / "docs" / "devops_collection_filling_guide.md").read_text(encoding="utf-8")
    manual = (ROOT / "docs" / "devops_dcp_collection_manual.md").read_text(encoding="utf-8")
    contract = (ROOT / "docs" / "dataflow_agent_input_contract_v0.1.md").read_text(encoding="utf-8")

    assert "DevOps" in collection_guide
    assert "docs/devops_dcp_collection_manual.md" in readme
    assert "scripts/setup_env.sh" in readme
    assert "workbook_schema.v0.1" in readme
    assert "dataflow_template.v1.0" in readme
    assert "samples/DCP_clean_v0.1/" in readme
    assert "docs/devops_dcp_collection_manual.md" in collection_guide
    for step in range(12):
        assert f"Step {step}" in manual
    for required in [
        "doctor.sh",
        "check_dcp.sh",
        "check_summary.md",
        "fix_list.md",
        "Evidence_ID",
        "Confirmation_Status",
        "Runtime_Type",
        "Target_Type",
        "Interaction_Mode",
        "templates/dataflow_v1.0/",
        "samples/DCP_clean_v0.1/",
        "docs/dataflow_agent_input_contract_v0.1.md",
        "schemas/workbook_schema.json",
    ]:
        assert required in manual
    assert "规则驱动确定性 Agent" in readme
    assert "不凭空补依赖" in readme
    assert "不自动接受安全例外" in readme
    assert "不修改生产环境" in readme
    assert "rule-driven deterministic agent" in contract
    assert "does not take over production environments" in contract


def test_changelog_tracks_current_and_historical_changes() -> None:
    changelog = (ROOT / "CHANGELOG.md").read_text(encoding="utf-8")
    pyproject = (ROOT / "pyproject.toml").read_text(encoding="utf-8")

    assert "# 中文版本" in changelog
    assert "# English Version" in changelog
    assert changelog.index("# 中文版本") < changelog.index("# English Version")
    assert "## Unreleased" in changelog
    assert "## 0.1.1 - 2026-07-01" in changelog
    assert "## 0.1.0 - 2026-06-24" in changelog
    assert 'version = "0.1.1"' in pyproject
    assert "0.1.1" in changelog
    assert "0.1.0" in changelog
    assert "历史开发记录 - 2026-06-23 至 2026-06-24" in changelog
    assert "Historical Development Record - 2026-06-23 To 2026-06-24" in changelog
    for required in [
        "MIT License",
        "service drilldown",
        "query-port",
        "Dropped graph edges",
        "field-level conflict diff",
        "DevOps DCP Collection Manual",
        "query_service_ports.sh",
        "doctor.sh",
        "setup_env.sh",
        "DCP_clean_v0.1",
        "executive overview readiness",
        "9dadfab",
        "f064829",
    ]:
        assert required in changelog


def test_license_and_generic_naming_are_enforced() -> None:
    license_text = (ROOT / "LICENSE").read_text(encoding="utf-8")
    assert license_text.startswith("MIT License")
    assert "Copyright (c) 2026 edmund-xl" in license_text
    assert "Permission is hereby granted, free of charge" in license_text
    assert 'license = "MIT"' in (ROOT / "pyproject.toml").read_text(encoding="utf-8")
    forbidden_license_terms = [
        "Pro" + "prietary",
        "pro" + "prietary",
        "all rights" + " are reserved",
        "No permission" + " is granted",
    ]
    for path in [
        ROOT / "README.md",
        ROOT / "CHANGELOG.md",
        ROOT / "LICENSE",
        ROOT / "pyproject.toml",
        ROOT / "src" / "dataflow_agent" / "packager.py",
        ROOT / "src" / "dataflow_agent" / "report_generator.py",
        ROOT / "src" / "dataflow_agent" / "summaries.py",
        ROOT / "docs" / "collector_quick_check_guide.md",
        ROOT / "docs" / "aggregation_operator_guide.md",
        ROOT / "docs" / "devops_dcp_collection_manual.md",
        ROOT / "docs" / "dataflow_agent_input_contract_v0.1.md",
    ]:
        text = path.read_text(encoding="utf-8")
        for term in forbidden_license_terms:
            assert term not in text
    legacy = "mega" + "eth"
    text_paths = [
        path
        for path in ROOT.rglob("*")
        if path.is_file()
        and ".git" not in path.parts
        and "dist" not in path.parts
        and "agent_check" not in path.parts
        and "__pycache__" not in path.parts
        and path.suffix.lower() in {".md", ".py", ".toml", ".sh", ".txt"}
    ]
    for path in text_paths:
        assert legacy not in path.name.lower()
        assert legacy not in path.read_text(encoding="utf-8").lower()

    for workbook_path in [SAMPLE_WORKBOOK, CLEAN_SAMPLE_WORKBOOK]:
        with ZipFile(workbook_path) as workbook:
            for name in workbook.namelist():
                if name.endswith((".xml", ".rels")):
                    data = workbook.read(name).decode("utf-8", errors="ignore").lower()
                    assert legacy not in data


def test_template_package_has_no_legacy_project_name() -> None:
    legacy_patterns = ["mega" + "eth", "mega" + " eth", "mega" + "-eth"]
    template_root = ROOT / "templates" / "dataflow_v1.0"
    expected_files = {
        "README.md",
        "dataflow_agent_io_contract_v1.0.md",
        "dataflow_collection_filling_guide_v1.0.docx",
        "dataflow_collection_template_bundle_v1.0.zip",
        "dataflow_data_dictionary_v1.0.xlsx",
        "dataflow_main_collection_template_v1.0.xlsx",
        "dataflow_overview_demo_v1.0.png",
        "dataflow_project_final_plan_v1.0.docx",
        "dataflow_sample_input_v1.0.xlsx",
        "dataflow_service_dependency_drilldown_demo_v1.0.png",
        "dataflow_task_collection_mapping_v1.0.xlsx",
    }

    assert template_root.exists()
    assert {path.name for path in template_root.iterdir() if path.is_file()} == expected_files

    def assert_no_legacy(text: str, label: str) -> None:
        lowered = text.lower()
        for pattern in legacy_patterns:
            assert pattern not in lowered, f"{label} contains legacy project name pattern {pattern!r}"

    def scan_zip_bytes(data: bytes, label: str) -> None:
        with ZipFile(BytesIO(data)) as archive:
            for name in archive.namelist():
                assert_no_legacy(name, f"{label}:{name}")
                payload = archive.read(name)
                lower_name = name.lower()
                if lower_name.endswith((".docx", ".xlsx", ".zip")):
                    scan_zip_bytes(payload, f"{label}:{name}")
                elif lower_name.endswith((".xml", ".rels", ".md", ".txt", ".json", ".csv")):
                    assert_no_legacy(payload.decode("utf-8", errors="ignore"), f"{label}:{name}")

    for path in template_root.rglob("*"):
        assert_no_legacy(str(path.relative_to(ROOT)), str(path.relative_to(ROOT)))
        if not path.is_file():
            continue
        lower_name = path.name.lower()
        if lower_name.endswith((".docx", ".xlsx", ".zip")):
            scan_zip_bytes(path.read_bytes(), str(path.relative_to(ROOT)))
        elif lower_name.endswith((".md", ".txt", ".json", ".csv")):
            assert_no_legacy(path.read_text(encoding="utf-8"), str(path.relative_to(ROOT)))


def test_template_package_exposes_runtime_and_target_fields() -> None:
    template_root = ROOT / "templates" / "dataflow_v1.0"
    main_template = load_workbook(template_root / "dataflow_main_collection_template_v1.0.xlsx", data_only=True)
    sample_template = load_workbook(template_root / "dataflow_sample_input_v1.0.xlsx", data_only=True)
    dictionary = load_workbook(template_root / "dataflow_data_dictionary_v1.0.xlsx", data_only=True)

    service_headers = [cell.value for cell in main_template["04_服务"][5]]
    dependency_headers = [cell.value for cell in main_template["05_依赖关系"][5]]
    for field in ["运行类型(Runtime_Type)", "运行ID(Runtime_ID)", "运行名称(Runtime_Name)", "运行命名空间(Runtime_Namespace)", "运行集群(Runtime_Cluster)", "运行区域(Runtime_Region)"]:
        assert field in service_headers
    for field in ["目标类型(Target_Type)", "目标ID(Target_ID)", "交互模式(Interaction_Mode)"]:
        assert field in dependency_headers

    sample_service_headers = [cell.value for cell in sample_template["04_服务"][5]]
    runtime_type_col = sample_service_headers.index("运行类型(Runtime_Type)") + 1
    assert sample_template["04_服务"].cell(7, runtime_type_col).value == "Kubernetes"

    dictionary_fields = {row[1] for row in dictionary["数据字典"].iter_rows(min_row=2, values_only=True)}
    assert "运行类型(Runtime_Type)" in dictionary_fields
    assert "目标类型(Target_Type)" in dictionary_fields

    with ZipFile(template_root / "dataflow_collection_template_bundle_v1.0.zip") as archive:
        nested_main = load_workbook(BytesIO(archive.read("dataflow_main_collection_template_v1.0.xlsx")), data_only=True)
        nested_headers = [cell.value for cell in nested_main["04_服务"][5]]
        assert "运行类型(Runtime_Type)" in nested_headers
        assert "目标类型(Target_Type)" in archive.read("README.md").decode("utf-8")
